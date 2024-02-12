import docx
import re
import os
import time
from tkinter import messagebox
from vuln_descriptions import vuln_info


def remove_html_tags(text):
    if not isinstance(text, str):
        text = str(text) if text is not None else ''
    return re.sub(r'<[^>]+>', '', text)

def clean_string(input_string):
    # Ensure the input is a string to avoid re.sub errors
    input_string = str(input_string) if input_string is not None else ''
    return re.sub(r'[\x00-\x1f\x7f-\x9f]', '', input_string)

# Other functions should ensure that any data passed into clean_string
# or similar operations is appropriately type-checked or converted to strings.


def log_error(e, report_name):
    timestamp = time.strftime('%Y-%m-%d %H-%M-%S')
    error_message = f"Error: {e}"
    log_file_name = f"{report_name}_error_log_{timestamp}.txt"
    log_file_path = os.path.join(os.getcwd(), log_file_name)
    with open(log_file_path, "a") as file:
        file.write(error_message + "\n")

def determine_vuln_level(cvss_score):
    if 0.0 <= cvss_score <= 3.9:
        return "低危"
    elif 4.0 <= cvss_score <= 6.9:
        return "中危"
    elif 7.0 <= cvss_score <= 8.9:
        return "高危"
    elif 9.0 <= cvss_score <= 10.0:
        return "严重"
    else:
        return "未知"

def generate_word_report(all_vulns, report_name, excluded_vulns=[]):
    print("开始创建 Word 文档")
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '黑体'

    # 聚合漏洞信息
    vuln_aggregate = {}
    for vuln_id, vuln in all_vulns.items():
        chinese_vt_name = vuln["chinese_vt_name"]
        english_vt_name = vuln["vt_name"]  # 假设原始英文名称在vt_name字段
        if chinese_vt_name not in excluded_vulns:
            # 优先使用中文名称，如果没有则使用英文名称
            key = chinese_vt_name if chinese_vt_name else english_vt_name
            if key not in vuln_aggregate:
                vuln_aggregate[key] = {"vulns": []}
            vuln_aggregate[key]["vulns"].append(vuln)

    # 创建漏洞统计表
    doc.add_heading("漏洞统计表", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '漏洞名称'
    hdr_cells[1].text = '影响URL'
    hdr_cells[2].text = '危险等级'
    for key, data in vuln_aggregate.items():
        urls = set()
        for vuln in data["vulns"]:
            url = vuln.get("combined_url", vuln["affects_url"])
            urls.add(url)
        row_cells = table.add_row().cells
        row_cells[0].text = key  # 可以选择使用中文或英文名称
        row_cells[1].text = '\n'.join(urls)
        # 假设所有同类型漏洞的危险等级相同
        row_cells[2].text = determine_vuln_level(data["vulns"][0].get("cvss_score", 0))

    # 为每种漏洞类型添加详细信息
    for key, data in vuln_aggregate.items():
        doc.add_heading(f"{key} 漏洞详情", level=1)
        for vuln in data["vulns"]:
            # 获取漏洞详细信息
            vuln_details = vuln_info.get(key, {})
            long_description = remove_html_tags(vuln_details.get("long_description", ""))
            affects_detail = remove_html_tags(vuln_details.get("affects_detail", ""))
            description = remove_html_tags(vuln_details.get("description", ""))
            impact = remove_html_tags(vuln_details.get("impact", ""))
            recommendation = remove_html_tags(vuln_details.get("recommendation", ""))
            details = vuln.get("details", "")
            request = vuln.get("request", "").replace('\\n', '').replace('\\r', '')
            combined_url = vuln.get("combined_url", vuln["affects_url"])
            level = determine_vuln_level(vuln.get("cvss_score", 0))


            # 添加漏洞信息到文档
            doc.add_heading(key, level=2)
            # doc.add_heading(english_vt_name, level=2)
            doc.add_heading("漏洞详情：", level=3)
            doc.add_paragraph(clean_string(f"{long_description}\n{affects_detail}\n{description}"))
            doc.add_heading("漏洞危害：", level=3)
            doc.add_paragraph(impact)
            doc.add_heading("漏洞URL：", level=3)
            doc.add_paragraph(combined_url)
            doc.add_heading("漏洞细节：", level=3)
            doc.add_paragraph(details)
            doc.add_heading("漏洞请求：", level=3)
            doc.add_paragraph(request)
            doc.add_heading("漏洞等级：", level=3)
            doc.add_paragraph(level)
            doc.add_heading("漏洞建议：", level=3)
            doc.add_paragraph(recommendation)
            doc.add_page_break()

    # 保存文档
    file_name = f"{clean_string(report_name)}_{time.strftime('%Y-%m-%d %H-%M-%S')}.docx"
    doc.save(file_name)
    print("已保存 Word 文档")
    messagebox.showinfo("操作完成", f"报告已生成: {file_name}")