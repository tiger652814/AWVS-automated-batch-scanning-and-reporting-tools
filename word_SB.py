#word_SB.py
import docx
import re
import os
import time
from tkinter import messagebox

# 假设这是从某处导入的漏洞描述信息
from vuln_descriptions import vuln_info
from special_vuln_descriptions import special_vuln_info

def remove_html_tags(text):
    if not isinstance(text, str):
        text = str(text) if text is not None else ''
    return re.sub(r'<[^>]+>', '', text)

def log_error(e, report_name):
    timestamp = time.strftime('%Y-%m-%d %H-%M-%S')
    error_message = f"Error: {e}"
    log_file_name = f"{report_name}_error_log_{timestamp}.txt"
    log_file_path = os.path.join(os.getcwd(), log_file_name)
    with open(log_file_path, "a") as file:
        file.write(error_message + "\n")

def remove_html_tags(text):
    """Safely remove HTML tags from a string."""
    # Convert non-string input to an empty string
    text = str(text) if text is not None else ''
    return re.sub(r'<[^>]+>', '', text)

def clean_string(input_string):
    """Safely clean unwanted characters from a string."""
    # Convert non-string input to an empty string
    input_string = str(input_string) if input_string is not None else ''
    return re.sub(r'[\x00-\x1f\x7f-\x9f]', '', input_string)

def determine_vuln_level(cvss_score):
    if cvss_score >= 0.0 and cvss_score <= 3.9:
        return "低危"
    elif cvss_score >= 4.0 and cvss_score <= 6.9:
        return "中危"
    elif cvss_score >= 7.0 and cvss_score <= 8.9:
        return "高危"
    elif cvss_score >= 9.0 and cvss_score <= 10.0:
        return "严重"
    else:
        return "未知"
def add_vuln_details_to_doc(doc, vuln, aggregated_urls):
    chinese_vt_name = vuln["chinese_vt_name"]
    english_vt_name = vuln["vt_name"]
    original_vt_name = english_vt_name  # 保存原始的漏洞名称

    # 如果存在中文标题则使用中文标题，否则使用英文标题
    display_vt_name = chinese_vt_name if chinese_vt_name else english_vt_name
    print(f"vdisplay_vt_name:{display_vt_name}")
    print(f"english_vt_name:{english_vt_name}")
    print(f"chinese_vt_name:{chinese_vt_name}")
    # 给 vuln_details 赋初始值
    vuln_details = vuln

    level = determine_vuln_level(vuln.get("cvss_score", 0))

    # 特殊标题处理的标志
    special_title_handled = False

    # 特殊标题处理，解析库名和版本
    if '易受攻击的 JavaScript 库' in display_vt_name or '已过时的 JavaScript 库' in display_vt_name:
        details = vuln.get("details", "")
        library_version_matches = re.findall(r'<strong>(.*?)</strong>', details)
        if library_version_matches:
            for match in library_version_matches:
                split_result = match.split()
                if len(split_result) >= 2:
                    library = ' '.join(split_result[:-1])
                    version = split_result[-1]
                    display_vt_name = f"{library} {version} - {('已过时的 JavaScript 库' if '已过时的 JavaScript 库' in display_vt_name else '易受攻击的 JavaScript 库')}"
                    # 处理库名和版本
                    # 这里可以添加您的逻辑来处理库名和版本
                    special_title_handled = True
                else:
                    display_vt_name = original_vt_name
                    special_title_handled = True
        else:
            display_vt_name = original_vt_name
            special_title_handled = True
    else:
        display_vt_name = original_vt_name

    # 如果已经处理了特殊标题，则不再修改 display_vt_name
    if not special_title_handled:
        display_vt_name = chinese_vt_name if chinese_vt_name else english_vt_name

    # 根据漏洞类型从相应的信息源提取信息
    is_outdated_js_lib = '已过时的 JavaScript 库' in display_vt_name
    is_vulnerable_js_lib = '易受攻击的 JavaScript 库' in display_vt_name

    if is_outdated_js_lib:
        vuln_details_key = "已过时的 JavaScript 库"
    elif is_vulnerable_js_lib:
        vuln_details_key = "易受攻击的 JavaScript 库"
    else:
        vuln_details_key = display_vt_name

    vuln_details = special_vuln_info.get(vuln_details_key,
                                         vuln_info.get(vuln_details_key, {"long_description": "信息不可用"}))

    # 移除HTML标签并处理数据
    long_description = remove_html_tags(vuln_details.get("long_description", ""))
    affects_detail = remove_html_tags(vuln_details.get("affects_detail", ""))
    description = remove_html_tags(vuln_details.get("description", ""))
    impact = remove_html_tags(vuln_details.get("impact", ""))
    recommendation = remove_html_tags(vuln_details.get("recommendation", ""))

    # 从原始 vuln 对象中获取 details 和 request
    # details = vuln.get("details", "")
    # request = vuln.get("request", "").replace('\\n', '').replace('\\r', '')
    details = str(vuln.get("details", ""))
    request = str(vuln.get("request", "")).replace('\\n', '').replace('\\r', '')

    # 清除漏洞细节中的HTML标签
    details_cleaned = remove_html_tags(details)

    # 将漏洞详细信息添加到文档中
    doc.add_heading(f"{display_vt_name} ", level=2)
    # doc.add_heading(f"{english_vt_name}", level=2)
    doc.add_heading("漏洞详情：", level=3)
    doc.add_paragraph(clean_string(f"{long_description}\n{affects_detail}\n{description}"))
    doc.add_heading("漏洞危害：", level=3)
    doc.add_paragraph(impact)

    # 处理并添加URLs
    # 将单个字符串拆分为URL列表
    aggregated_urls = aggregated_urls.split(',')

    # 使用换行符连接URL列表
    urls_string = '\n'.join(aggregated_urls)

    # 添加到文档
    doc.add_heading("漏洞URL：", level=3)
    doc.add_paragraph(urls_string)

    doc.add_heading("漏洞细节：", level=3)
    doc.add_paragraph(details_cleaned)
    doc.add_heading("漏洞请求：", level=3)
    doc.add_paragraph(request)
    doc.add_heading("漏洞等级：", level=3)
    doc.add_paragraph(str(level))
    doc.add_heading("漏洞建议：", level=3)
    doc.add_paragraph(recommendation)

    return display_vt_name
def generate_word_report(all_vulns, report_name, excluded_vulns=[]):
    print("开始创建 Word 文档")
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '黑体'

    # 用于存储不同危险等级的漏洞数量
    vuln_level_counts = {"低危": 0, "中危": 0, "高危": 0, "严重": 0, "未知": 0}

    # 用于存储相同漏洞名称的 URL 和其他信息
    vuln_aggregate = {}

    try:
        # 添加漏洞统计表
        doc.add_heading("漏洞统计表", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '漏洞名称'
        hdr_cells[1].text = '影响URL'
        hdr_cells[2].text = '危险等级'

        for vuln_id, vuln in all_vulns.items():
            chinese_vt_name = vuln["chinese_vt_name"]
            english_vt_name = vuln["vt_name"]
            display_vt_name = chinese_vt_name if chinese_vt_name else english_vt_name
            level = determine_vuln_level(vuln["cvss_score"])
            key = display_vt_name
            if chinese_vt_name not in excluded_vulns:
                if key not in vuln_aggregate:
                    vuln_aggregate[key] = {"vuln": vuln, "urls": set()}
                vuln_aggregate[key]["urls"].add(vuln.get("combined_url", vuln["affects_url"]))

                vuln_level_counts[level] += 1

        # 填充漏洞统计表
        for display_vt_name, data in vuln_aggregate.items():
            vuln = data["vuln"]
            row_cells = table.add_row().cells
            row_cells[0].text = display_vt_name
            row_cells[1].text = '\n'.join(data["urls"])
            row_cells[2].text = determine_vuln_level(vuln["cvss_score"])

        # 添加漏洞等级统计表
        doc.add_heading("漏洞等级统计", level=2)
        level_table = doc.add_table(rows=1, cols=2)
        level_table.style = 'Table Grid'
        hdr_cells = level_table.rows[0].cells
        hdr_cells[0].text = '危险等级'
        hdr_cells[1].text = '漏洞数量'

        # 填充漏洞等级统计表
        for level, count in vuln_level_counts.items():
            row_cells = level_table.add_row().cells
            row_cells[0].text = level
            row_cells[1].text = str(count)

        # 添加每个漏洞的详细信息
        for display_vt_name, data in vuln_aggregate.items():
            vuln = data["vuln"]
            urls = '\n'.join(data["urls"])
            add_vuln_details_to_doc(doc, vuln, urls)

        # 保存文档
        file_name = f"{clean_string(report_name)}_{time.strftime('%Y-%m-%d %H-%M-%S')}.docx"
        doc.save(file_name)
        print("已保存 Word 文档")
        messagebox.showinfo("操作完成", f"报告已生成: {file_name}")

    except Exception as e:
        print(f"生成 Word 文档时发生错误: {e}")
        log_error(e, report_name)